from pathlib import Path
from datetime import datetime

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\23178\JL\work\WT9011DCL_BT50_FW")
OUT_DIR = ROOT / "project_docs"
ASSET_DIR = OUT_DIR / "low_power_report_assets"
DOCX_PATH = OUT_DIR / "WT9011DCL-BT50低功耗代码开发报告.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(95, 95, 95)
FILL = "F2F4F7"
BORDER = "C9D3DF"


def font_path():
    for p in [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]:
        if p.exists():
            return str(p)
    return None


FONT_PATH = font_path()


def pil_font(size, bold=False):
    return ImageFont.truetype(FONT_PATH, size) if FONT_PATH else ImageFont.load_default()


def set_run_font(run, size=None, color=None, bold=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph(text="", style=None, bold=False, color=INK, size=11, after=6, before=0):
    p = DOC.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def heading(text, level=1):
    style = "Heading {}".format(level)
    p = DOC.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level],
                 color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_bullets(items):
    for item in items:
        p = DOC.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_numbered(items):
    for item in items:
        p = DOC.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_table(headers, rows, widths):
    table = DOC.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, FILL)
        set_cell_text(cell, h, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], value, size=9.5)
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def draw_rounded_box(draw, xy, text, fill, outline, font, align="center"):
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    line_heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(line_heights) + (len(lines) - 1) * 8
    y = y1 + (y2 - y1 - total_h) / 2
    for line, lh in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        if align == "center":
            x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        else:
            x = x1 + 18
        draw.text((x, y), line, fill=(30, 45, 60), font=font)
        y += lh + 8


def arrow(draw, start, end, color=(57, 85, 120)):
    draw.line([start, end], fill=color, width=3)
    ex, ey = end
    sx, sy = start
    if ex > sx:
        pts = [(ex, ey), (ex - 12, ey - 7), (ex - 12, ey + 7)]
    elif ex < sx:
        pts = [(ex, ey), (ex + 12, ey - 7), (ex + 12, ey + 7)]
    else:
        pts = [(ex, ey), (ex - 7, ey - 12), (ex + 7, ey - 12)]
    draw.polygon(pts, fill=color)


def make_chain_image(path):
    img = Image.new("RGB", (1500, 760), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(34)
    font = pil_font(24)
    small = pil_font(20)
    d.text((50, 36), "SDK低功耗调用链", fill=(20, 55, 90), font=title_font)
    boxes = [
        ((60, 150, 285, 260), "板级配置\nTCFG_LOWPOWER"),
        ((355, 150, 600, 260), "board_power_init()\npower_init()"),
        ((670, 150, 930, 260), "REGISTER_LP_TARGET\n注册idle目标"),
        ((1000, 150, 1255, 260), "系统空闲扫描\n逐个调用is_idle()"),
    ]
    for box, text in boxes:
        draw_rounded_box(d, box, text, (232, 238, 245), (150, 172, 198), font)
    for a, b in [((285, 205), (355, 205)), ((600, 205), (670, 205)), ((930, 205), (1000, 205))]:
        arrow(d, a, b)
    draw_rounded_box(d, (385, 390, 650, 510), "全部idle=1\n允许SDK进入powerdown", (225, 244, 232), (95, 160, 110), font)
    draw_rounded_box(d, (850, 390, 1115, 510), "任意idle=0\n保持运行不睡眠", (252, 238, 218), (185, 135, 65), font)
    arrow(d, (1125, 260), (985, 390))
    arrow(d, (1125, 260), (535, 390))
    draw_rounded_box(d, (1190, 390, 1430, 510), "sleep_enter/exit\n板级回调", (242, 242, 242), (160, 160, 160), font)
    arrow(d, (650, 450), (1190, 450))
    d.text((60, 610), "本阶段只增加WTYI的idle查询门控，不做Poweroff、不改唤醒脚、不调BLE连接参数。",
           fill=(90, 90, 90), font=small)
    img.save(path)


def make_gate_image(path):
    img = Image.new("RGB", (1500, 820), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(34)
    font = pil_font(24)
    small = pil_font(20)
    d.text((50, 34), "WTYI低功耗busy门控", fill=(20, 55, 90), font=title_font)
    flags = [
        ("BT busy", 95), ("SPI busy", 190), ("IIC busy", 285),
        ("ADC busy", 380), ("OTA busy", 475)
    ]
    for text, y in flags:
        draw_rounded_box(d, (80, y, 310, y + 62), text, (245, 247, 250), (180, 190, 205), font)
        arrow(d, (310, y + 31), (550, 395))
    draw_rounded_box(d, (550, 315, 845, 475), "wtyi_power_busy_mask\n五个标志汇总", (232, 238, 245), (120, 150, 185), font)
    arrow(d, (845, 395), (980, 395))
    draw_rounded_box(d, (980, 315, 1265, 475), "wtyi_power_idle_query()\nmask==0 ?", (232, 238, 245), (120, 150, 185), font)
    draw_rounded_box(d, (1030, 575, 1310, 675), "mask=0\n返回1: 可睡", (225, 244, 232), (95, 160, 110), font)
    draw_rounded_box(d, (620, 575, 900, 675), "mask!=0\n返回0: 禁止睡", (252, 238, 218), (185, 135, 65), font)
    arrow(d, (1125, 475), (1170, 575))
    arrow(d, (1065, 475), (760, 575))
    d.text((80, 735), "OTA开始后始终置busy，OTA结束后清除；SPI/IIC/ADC只在读写窗口置busy。",
           fill=(90, 90, 90), font=small)
    img.save(path)


def make_code_image(path):
    code = [
        "static u8 wtyi_power_idle_query(void)",
        "{",
        "#if WTYI_LOW_POWER_ENABLE",
        "    return wtyi_power_busy_mask == 0;",
        "#else",
        "    return 0;",
        "#endif",
        "}",
        "",
        "REGISTER_LP_TARGET(wtyi_lp_target) = {",
        "    .name = \"wtyi\",",
        "    .is_idle = wtyi_power_idle_query,",
        "};",
    ]
    img = Image.new("RGB", (1500, 620), (248, 250, 252))
    d = ImageDraw.Draw(img)
    title_font = pil_font(30)
    mono = ImageFont.truetype(str(Path(r"C:\Windows\Fonts\consola.ttf")), 24) if Path(r"C:\Windows\Fonts\consola.ttf").exists() else pil_font(24)
    d.rounded_rectangle((45, 40, 1455, 570), radius=18, fill=(35, 45, 58), outline=(100, 120, 140), width=2)
    d.text((85, 75), "关键代码摘录：注册WTYI低功耗idle查询目标", fill=(235, 240, 246), font=title_font)
    y = 140
    for line in code:
        d.text((95, y), line, fill=(230, 238, 247), font=mono)
        y += 32
    img.save(path)


def add_figure(path, caption):
    p = DOC.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    cap = DOC.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, color=MUTED, bold=True)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("WT9011DCL-BT50 | 低功耗代码开发报告")
    set_run_font(r, size=9.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Generated: {}".format(datetime.now().strftime("%Y-%m-%d %H:%M")))
    set_run_font(r, size=9, color=MUTED)


ASSET_DIR.mkdir(parents=True, exist_ok=True)
make_chain_image(ASSET_DIR / "sdk_low_power_chain.png")
make_gate_image(ASSET_DIR / "wtyi_busy_gate.png")
make_code_image(ASSET_DIR / "wtyi_code_snippet.png")

DOC = Document()
configure_doc(DOC)

p = DOC.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("代码开发报告")
set_run_font(r, size=12, color=MUTED, bold=True)

p = DOC.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("WT9011DCL-BT50 低功耗基础框架")
set_run_font(r, size=24, color=INK, bold=True)

p = DOC.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("项目：WT9011DCL-BT50 | 主控：AC6321A4 | SDK：fw-AC63_BT_SDK")
set_run_font(r, size=11, color=MUTED)

add_table(
    ["项目", "内容"],
    [
        ["本阶段目标", "建立最基础、可理解、可回退的低功耗门控框架"],
        ["完成状态", "代码已开发，离线编译通过，未执行烧录，未声明真实功耗"],
        ["提交号", "8187461 添加WTYI低功耗门控框架"],
        ["最终固件", r"D:\23178\JL\work\WT9011DCL_BT50_FW\cpu\bd19\tools\app.bin"],
        ["SHA256", "74DF7FDB96B318E62114782CDD2361161C0AAE3F5AD9D5C9B02BD8208ED90221"],
    ],
    [1.45, 5.05],
)

heading("1. 开发背景和边界", 1)
paragraph("本阶段面向低功耗初学者，不追求最低功耗，也不做复杂参数优化。目标是先把“什么时候允许睡眠、什么时候禁止睡眠”的软件框架搭起来，保证代码结构清楚、可开关、可回退。")
add_bullets([
    "不进入Poweroff，不修改软关机流程。",
    "不修改唤醒引脚，不改变硬件接线要求。",
    "不修改BLE连接参数，避免引入蓝牙连接稳定性风险。",
    "不执行烧录，不声称已经测得真实功耗。",
])

heading("2. SDK现有低功耗调用链", 1)
paragraph("杰理SDK的思路不是由业务代码直接让芯片睡眠，而是每个模块告诉系统自己是否空闲。系统在空闲任务里统一判断，只有全部模块都允许睡眠时，才会进入sleep或powerdown。")
add_figure(ASSET_DIR / "sdk_low_power_chain.png", "图1 SDK低功耗调用链")

heading("3. 当前工程低功耗配置", 1)
add_table(
    ["配置项", "所在文件", "当前值/说明"],
    [
        ["TCFG_LOWPOWER_LOWPOWER_SEL", "board_wt9011dcl_bt50_cfg.h", "SLEEP_EN，当前板级工程已开启powerdown能力"],
        ["TCFG_LOWPOWER_POWER_SEL", "board_wt9011dcl_bt50_cfg.h", "PWR_DCDC15，本阶段未修改"],
        ["sleep_enter_callback", "board_wt9011dcl_bt50.c", "SDK进入sleep前调用，本阶段未改回调逻辑"],
        ["sleep_exit_callback", "board_wt9011dcl_bt50.c", "SDK退出sleep后调用，本阶段未改回调逻辑"],
        ["REGISTER_LP_TARGET", "power_interface.h", "SDK低功耗目标注册宏，本阶段新增wtyi目标"],
    ],
    [2.35, 1.95, 2.20],
)

heading("4. 本阶段代码开发步骤", 1)
add_numbered([
    "搜索SDK中的REGISTER_LP_TARGET、低功耗宏、sleep_enter_callback和sleep_exit_callback，确认调用方式。",
    "确认WT9011DCL-BT50板级工程已经通过TCFG_LOWPOWER_LOWPOWER_SEL启用低功耗能力。",
    "新增wtyi_power_manager.c和wtyi_power_manager.h，独立管理WTYI低功耗门控逻辑。",
    "注册wtyi_lp_target，让SDK能够查询WTYI模块当前是否空闲。",
    "建立BT、SPI、IIC、ADC、OTA五个busy标志，并用bit mask统一管理。",
    "接入OTA状态，OTA开始置busy，OTA结束清busy，防止升级期间进入低功耗。",
    "接入SPI、IIC、ADC测试驱动，在初始化和读写采样期间置busy，操作结束后清busy。",
    "将新源文件加入Makefile和CodeBlocks工程文件，保证命令行和CodeBlocks都能编译。",
    "执行build_only.bat完成离线编译，确认未调用下载脚本和烧录工具。",
    "生成low_power_design.md和本Word报告，记录设计说明、回退方法和后续验收动作。",
])

heading("5. WTYI低功耗门控设计", 1)
paragraph("WTYI模块新增一个全局busy mask。BT、SPI、IIC、ADC、OTA分别占用一个bit。任意bit为1时，WTYI向SDK报告“忙”，SDK不进入低功耗；所有bit为0时，WTYI向SDK报告“空闲”，是否真正进入powerdown仍由SDK和其他模块共同决定。")
add_figure(ASSET_DIR / "wtyi_busy_gate.png", "图2 WTYI busy门控逻辑")

heading("6. 关键代码说明", 1)
add_figure(ASSET_DIR / "wtyi_code_snippet.png", "图3 wtyi_power_manager.c关键代码摘录")
add_table(
    ["代码点", "作用"],
    [
        ["wtyi_power_busy_mask", "保存BT/SPI/IIC/ADC/OTA五类busy状态"],
        ["wtyi_power_idle_query()", "SDK低功耗扫描时调用，返回1表示WTYI空闲，返回0表示禁止睡眠"],
        ["REGISTER_LP_TARGET(wtyi_lp_target)", "把WTYI模块注册进SDK低功耗目标表"],
        ["WTYI_LOW_POWER_ENABLE", "一键开关；为0时WTYI始终返回busy，阻止低功耗"],
        ["wtyi_power_ota_begin/end()", "OTA期间禁止低功耗，OTA结束后恢复允许判断"],
    ],
    [2.2, 4.3],
)

heading("7. 修改文件清单", 1)
add_table(
    ["文件", "修改内容"],
    [
        ["apps/spp_and_le/wtyi/wtyi_power_manager.c", "新增WTYI低功耗busy门控实现"],
        ["apps/spp_and_le/wtyi/wtyi_power_manager.h", "新增busy枚举和对外接口"],
        ["apps/spp_and_le/wtyi/wtyi_config.h", "新增WTYI_LOW_POWER_ENABLE宏"],
        ["apps/spp_and_le/wtyi/wtyi_board_test.c", "初始化WTYI低功耗管理模块"],
        ["apps/common/update/update.c", "OTA开始/结束同步设置WTYI OTA busy"],
        ["wtyi_spi_imu.c / wtyi_iic_mag.c / wtyi_adc_battery.c", "外设操作窗口置busy，结束清busy"],
        ["apps/spp_and_le/board/bd19/Makefile", "加入wtyi_power_manager.c编译"],
        ["apps/spp_and_le/board/bd19/AC632N_spp_and_le.cbp", "加入CodeBlocks工程文件列表"],
        ["project_docs/low_power_design.md", "记录设计说明、编译结果和回退方法"],
    ],
    [3.0, 3.5],
)

heading("8. 编译验证结果", 1)
add_table(
    ["项目", "结果"],
    [
        ["编译命令", r"D:\23178\JL\work\WT9011DCL_BT50_FW\tools\wtyi\build_only.bat"],
        ["编译结果", "BUILD OK，0个编译错误"],
        ["烧录状态", "未执行烧录，未运行download.bat、isd_download.exe"],
        ["固件路径", r"D:\23178\JL\work\WT9011DCL_BT50_FW\cpu\bd19\tools\app.bin"],
        ["固件大小", "207472 bytes"],
        ["固件SHA256", "74DF7FDB96B318E62114782CDD2361161C0AAE3F5AD9D5C9B02BD8208ED90221"],
        ["编译日志", r"D:\23178\JL\work\WT9011DCL_BT50_FW\build_logs\build_only_20260727_211606.log"],
    ],
    [1.6, 4.9],
)
paragraph("说明：链接阶段仍有SDK既有的LLVM gold plugin stack-size warning，本阶段没有新增编译或链接错误。")

heading("9. 回退方法", 1)
paragraph("最快回退方式是在wtyi_config.h中把WTYI_LOW_POWER_ENABLE改为0，然后重新执行build_only.bat。这样不需要删除代码，WTYI低功耗目标会一直报告busy，SDK不会因为WTYI模块而进入低功耗。")
add_bullets([
    "快速关闭：WTYI_LOW_POWER_ENABLE = 0。",
    "完整回退：删除wtyi_power_manager.c/.h，并从Makefile、CodeBlocks工程和OTA接入点移除相关引用。",
    "回退后必须重新编译，确认无未实现函数和链接错误。",
])

heading("10. 后续硬件验收建议", 1)
add_bullets([
    "先烧录新固件，确认蓝牙广播和连接仍正常。",
    "如果UART日志可用，观察[WTYI_BOOT] low_power=enabled mask=0x0。",
    "确认没有OTA过程时，系统应允许SDK根据自身策略进入powerdown。",
    "进行SPI/IIC/ADC测试时，外设操作窗口内应保持busy，操作结束后清busy。",
    "最后再接电流表或功耗分析仪测量真实功耗，不用估算值替代实测值。",
])

heading("结论", 1)
paragraph("本阶段已经完成WT9011DCL-BT50的低功耗基础代码框架。代码只做可理解、可回退的idle门控，不进入Poweroff，不修改唤醒源，不调整BLE连接参数。离线编译已通过，下一步是在烧录链路稳定后进行蓝牙启动、日志和电流实测验证。")

DOC.save(DOCX_PATH)
print(DOCX_PATH)
